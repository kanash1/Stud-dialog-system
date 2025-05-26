import os
import logging
from io import BytesIO
from typing import Any, List, Tuple, AsyncGenerator

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.ext.asyncio import AsyncEngine

from langchain.chains import (
    create_history_aware_retriever, create_retrieval_chain
)
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_community.chat_models import ChatOpenAI
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_postgres.v2.vectorstores import AsyncPGVectorStore
from langchain_postgres.v2.engine import PGEngine
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.messages.ai import AIMessageChunk

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Исключение для неподдерживаемых форматов файлов
class UnsupportedFileFormatError(Exception):
    pass

# Исключение при ошибке чтения файла
class ReadFileError(Exception):
    pass

# Хранилище документов с возможностью векторизации
class DocumentStore:
    def __init__(self):
        self._vector_store = None
        self._engine = None

    # Инициализация хранилища документов с векторной базой
    async def init_document_store(self, engine: AsyncEngine):
        logger.info("Initializing document store...")

        embeddings = HuggingFaceEmbeddings(
            model_name="deepvk/USER2-small",
            encode_kwargs={"prompt_name": "search_document"},
            query_encode_kwargs={"prompt_name": "search_query"}
        )

        self._engine = PGEngine.from_engine(engine)

        self._vector_store = await AsyncPGVectorStore.create(
            engine=self._engine,
            embedding_service=embeddings,
            table_name='doc_collection'
        )
        logger.info("Document store initialized.")

    # Добавление документов в хранилище
    async def add_documents(
        self,
        files: List[Tuple[str, bytes]],
        engine: AsyncEngine | None = None
    ):
        if not self._vector_store or not self._engine:
            raise RuntimeError("Storage is not initialized.")

        engine = engine or self._engine

        docs = []
        for filename, content in files:
            logger.info(f"Extracting text from: {filename}")
            extracted_docs = await self._extract_text_from_document(
                filename, content
            )
            docs.extend(extracted_docs)

        await self._vector_store.aadd_documents(docs, engine=engine)
        logger.info("Documents added to vector store.")

    # Извлечение текста из различных типов документов
    async def _extract_text_from_document(
        self, filename: str, content: bytes
    ) -> List[Document]:
        extension = os.path.splitext(filename)[1].lower()
        document = None

        if extension == ".txt":
            text = content.decode("utf-8", errors="ignore")
            document = Document(
                page_content=text, metadata={"filename": filename}
            )

        elif extension == ".docx":
            try:
                docx = DocxDocument(BytesIO(content))
                full_text = "\n".join(p.text for p in docx.paragraphs)
                document = Document(
                    page_content=full_text,
                    metadata={"filename": filename}
                )
            except Exception as e:
                raise ReadFileError(f"DOCX read error: {e}")

        elif extension == ".pdf":
            try:
                reader = PdfReader(BytesIO(content))
                pages = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        pages.append(Document(
                            page_content=text,
                            metadata={"filename": filename, "page": i + 1}
                        ))
                return RecursiveCharacterTextSplitter(
                    chunk_size=512, chunk_overlap=256
                ).split_documents(pages)
            except Exception as e:
                raise ReadFileError(f"PDF read error: {e}")

        else:
            raise UnsupportedFileFormatError(
                f"Unsupported file format: {extension}"
            )

        return RecursiveCharacterTextSplitter(
            chunk_size=512, chunk_overlap=256
        ).split_documents([document])

    @property
    def vector_store(self):
        if not self._vector_store:
            raise RuntimeError("Storage is not initialized.")
        return self._vector_store

# Исключение при неизвестном адаптере LLM
class UnknownAdapterError(Exception):
    pass

# Основной класс RAG-системы
class RAGSystem:
    def __init__(self, model: str, adapters: List[str]):
        self._model = model
        self._adapters = adapters
        self._doc_store = None

    # Инициализация хранилища документов
    async def _init_document_store(self, engine: AsyncEngine):
        logger.info("Initializing document store inside RAG system...")
        self._doc_store = DocumentStore()
        await self._doc_store.init_document_store(engine)

    @property
    def document_store(self):
        if self._doc_store is None:
            raise RuntimeError("Storage is not initialized.")
        return self._doc_store

    # Создание LLM с нужным адаптером
    def _create_llm(self, adapter: str | None = None):
        if adapter is None:
            adapter = self._model
        elif adapter not in self._adapters:
            raise UnknownAdapterError("Unknown adapter")
        return ChatOpenAI(model=adapter, temperature=0.3)

    # Создание retriever, учитывающего историю
    def _create_history_aware_retriever(self, llm):
        if self._doc_store is None:
            raise RuntimeError("Storage is not initialized.")

        system_prompt = (
            "Given the chat history and the user's last question, "
            "reformulate the question so that it is "
            "understandable without the history. "
            "DON'T answer the question, just reformulate it."
        )

        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{input}")
        ])

        return create_history_aware_retriever(
            llm=llm,
            retriever=self._doc_store.vector_store.as_retriever(
                search_type="similarity_score_threshold",
                search_kwargs={"score_threshold": 0.65, "k": 15}
            ),
            prompt=prompt
        )

    # Создание цепочки на основе документов
    def _create_document_chain(self, llm):
        prompt = ChatPromptTemplate.from_messages([
            ("system", (
                "You're a question-answering assistant. "
                "Use the context only if relevant. If not, "
                "answer from general knowledge. "
                "If unsure, say so.\n\n{context}"
            )),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{input}")
        ])
        return create_stuff_documents_chain(llm=llm, prompt=prompt)

    # Создание цепочки без использования RAG
    def _create_non_rag_chain(self, llm):
        prompt = ChatPromptTemplate.from_messages([
            ("system", "You're a helpful assistant."),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{input}")
        ])
        return prompt | llm

    # Создание полной RAG-цепочки
    def _create_rag_chain(self, llm):
        retriever = self._create_history_aware_retriever(llm)
        document_chain = self._create_document_chain(llm)
        return create_retrieval_chain(
            retriever=retriever, combine_docs_chain=document_chain
        )

    # Публичная инициализация системы
    async def init_system(self, engine: AsyncEngine):
        await self._init_document_store(engine)
        logger.info("RAG system initialized.")

    # Основной метод вызова: с RAG или без
    async def invoke(
        self,
        input: dict[str, Any],
        use_rag: bool = False,
        adapter: str = None
    ) -> str:
        llm = self._create_llm(adapter)

        if use_rag:
            logger.info("Invoking RAG chain...")
            result = await self._create_rag_chain(llm).ainvoke(input)
            return result['answer']

        logger.info("Invoking non-RAG chain...")
        result = await self._create_non_rag_chain(llm).ainvoke(input)
        return result.text()
    
    async def stream_invoke(
        self,
        input: dict[str, Any],
        use_rag: bool = False,
        adapter: str = None
    ) -> AsyncGenerator[AIMessageChunk, None]:
        llm = self._create_llm(adapter)

        chain = (
            self._create_rag_chain(llm)
            if use_rag else
            self._create_non_rag_chain(llm)
        )

        async for chunk in chain.astream(input):
            yield chunk
